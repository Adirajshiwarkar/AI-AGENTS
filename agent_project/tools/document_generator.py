import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from utils.logger import logger, log_generating_docx

class DocumentGenerator:
    """Generates professional DOCX files from structured markdown-based text."""

    def __init__(self, output_dir: str = None):
        if output_dir is None:
            output_dir = "/tmp/generated_docs" if os.getenv("VERCEL") else "generated_docs"
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def _add_page_number(self, run):
        """Helper to inject page number XML field into footer."""
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

    def _apply_text_formatting(self, run, font_name="Arial", size_pt=11, color_rgb=(51, 51, 51), bold=False, italic=False):
        """Helper to apply basic font styles."""
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.color.rgb = RGBColor(*color_rgb)
        run.bold = bold
        run.italic = italic

    def generate_docx(self, title: str, sections_content: dict, output_filename: str) -> str:
        """Creates a professional docx file using structured contents."""
        log_generating_docx()
        doc = Document()
        
        # Page Margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            
            # Setup Footer with Page Numbers
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = footer_para.add_run("Page ")
            self._apply_text_formatting(run, font_name="Arial", size_pt=9, color_rgb=(120, 120, 120))
            self._add_page_number(run)

        # 1. Document Title / Cover Header
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(36)
        title_para.paragraph_format.space_after = Pt(24)
        title_run = title_para.add_run(title)
        self._apply_text_formatting(title_run, font_name="Arial", size_pt=26, color_rgb=(0, 51, 102), bold=True)

        # Subtitle
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_para.paragraph_format.space_after = Pt(48)
        sub_run = sub_para.add_run("Prepared Autonomously by AI Document Generator Engine")
        self._apply_text_formatting(sub_run, font_name="Arial", size_pt=12, color_rgb=(102, 102, 102), italic=True)

        doc.add_page_break()

        # 2. Iterate through sections
        for sec_title, markdown_content in sections_content.items():
            if not markdown_content:
                continue

            # Heading
            h_para = doc.add_heading(level=1)
            h_para.paragraph_format.keep_with_next = True
            h_para.paragraph_format.space_before = Pt(18)
            h_para.paragraph_format.space_after = Pt(8)
            h_run = h_para.add_run(sec_title)
            self._apply_text_formatting(h_run, font_name="Arial", size_pt=18, color_rgb=(0, 51, 102), bold=True)

            # Parse markdown body
            self._parse_markdown_to_docx(doc, markdown_content)
            
            # Separate sections by page break except the last one
            doc.add_paragraph().paragraph_format.space_after = Pt(12)

        file_path = os.path.join(self.output_dir, output_filename)
        doc.save(file_path)
        logger.info(f"Document successfully saved to {file_path}")
        return file_path

    def _parse_markdown_to_docx(self, doc, markdown: str):
        """Basic line-based Markdown parser to handle headings, bullets, tables, and paragraphs."""
        lines = [line.strip() for line in markdown.split("\n")]
        
        in_table = False
        table_rows = []
        
        for line in lines:
            if not line:
                # Flush table if we were building one
                if in_table and table_rows:
                    self._render_docx_table(doc, table_rows)
                    table_rows = []
                    in_table = False
                continue

            # Check if line is part of a markdown table
            if line.startswith("|"):
                in_table = True
                # Skip separating lines like |---|---|
                if not all(c in "| -:\t" for c in line):
                    table_rows.append(line)
                continue
            elif in_table:
                # Flush table if we encountered a non-table line
                if table_rows:
                    self._render_docx_table(doc, table_rows)
                    table_rows = []
                in_table = False

            # Check headings
            if line.startswith("### "):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)
                r = p.add_run(line[4:])
                self._apply_text_formatting(r, font_name="Arial", size_pt=13, color_rgb=(0, 80, 150), bold=True)
            elif line.startswith("#### "):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(line[5:])
                self._apply_text_formatting(r, font_name="Arial", size_pt=11, color_rgb=(51, 102, 153), bold=True)
            # Bullet Lists
            elif line.startswith("* ") or line.startswith("- "):
                clean_text = line[2:]
                p = doc.add_paragraph(style='List Bullet')
                p.paragraph_format.space_after = Pt(3)
                
                # Check for simple bold in markdown bullet point (**text**)
                self._add_formatted_text(p, clean_text)
            # Regular paragraphs
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.line_spacing = 1.15
                self._add_formatted_text(p, line)
                
        # Final flush for tables at the end of the markdown string
        if in_table and table_rows:
            self._render_docx_table(doc, table_rows)

    def _add_formatted_text(self, paragraph, text: str):
        """Splits text by markdown inline bold markers `**` and renders appropriately."""
        parts = text.split("**")
        is_bold = False
        for part in parts:
            if not part:
                is_bold = not is_bold
                continue
            r = paragraph.add_run(part)
            self._apply_text_formatting(r, font_name="Arial", size_pt=11, bold=is_bold)
            is_bold = not is_bold

    def _render_docx_table(self, doc, markdown_rows: list):
        """Parses list of markdown table rows and inserts a styled Word table."""
        if not markdown_rows:
            return

        parsed_rows = []
        for row in markdown_rows:
            # Split and clean empty strings at margins
            cells = [c.strip() for c in row.split("|")]
            if cells and cells[0] == "":
                cells.pop(0)
            if cells and cells[-1] == "":
                cells.pop()
            parsed_rows.append(cells)

        if not parsed_rows:
            return

        num_cols = len(parsed_rows[0])
        num_rows = len(parsed_rows)
        
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Style table using basic borders and background color headers
        for r_idx, row_cells in enumerate(parsed_rows):
            # Safe boundary check in case markdown table columns are uneven
            for c_idx in range(min(num_cols, len(row_cells))):
                cell = table.cell(r_idx, c_idx)
                cell.text = row_cells[c_idx]
                
                # Format text
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    for r in p.runs:
                        is_header = (r_idx == 0)
                        self._apply_text_formatting(
                            r, 
                            font_name="Arial", 
                            size_pt=10, 
                            color_rgb=(255, 255, 255) if is_header else (51, 51, 51),
                            bold=is_header
                        )
                
                # Apply background shading for header row using XML manipulation
                if r_idx == 0:
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), '003366') # Deep navy blue header background
                    tcPr.append(shd)
                    
        # Add a blank paragraph after the table
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
